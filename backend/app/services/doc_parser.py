"""
文档解析服务
- 从 DOC/DOCX/PDF 提取纯文本
- 调用豆包大模型结构化为 JSON
- 支持简历与职位两种解析模式
"""
import json
from typing import Any

from loguru import logger

from app.ai.ark_client import ark_client
from app.ai.prompts import build_job_messages, build_resume_messages


class DocParser:
    """文档解析器 (文本提取 + 灵犀结构化)"""

    def extract_text(self, file_path: str) -> str:
        """从文件提取纯文本"""
        lower = file_path.lower()
        if lower.endswith(".docx"):
            return self._extract_docx(file_path)
        if lower.endswith(".pdf"):
            return self._extract_pdf(file_path)
        if lower.endswith(".doc"):
            # 旧版 .doc 格式, 尝试用 pdfplumber 或提示转换
            logger.warning(f"旧版 .doc 格式支持有限: {file_path}, 尝试用 pdfplumber 解析")
            return self._extract_pdf(file_path)
        raise ValueError(f"不支持的文件格式: {file_path}")

    def _extract_docx(self, file_path: str) -> str:
        """从 DOCX 提取文本"""
        from docx import Document

        doc = Document(file_path)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        # 表格内容
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))
        text = "\n".join(paragraphs)
        logger.info(f"DOCX 文本提取完成: {len(text)} 字符")
        return text

    def _extract_pdf(self, file_path: str) -> str:
        """从 PDF 提取文本"""
        import pdfplumber

        texts: list[str] = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                if page_text.strip():
                    texts.append(page_text.strip())
        text = "\n".join(texts)
        logger.info(f"PDF 文本提取完成: {len(text)} 字符, {len(texts)} 页")
        return text

    def parse_resume(self, text: str) -> tuple[dict[str, Any], dict[str, int]]:
        """用灵犀大模型将简历文本结构化, 返回 (结构化数据, token_usage)"""
        if not text.strip():
            raise ValueError("简历文本为空")
        # 截断超长文本 (避免 token 超限)
        truncated = text[:8000]
        messages = build_resume_messages(truncated)
        # 使用 mini 轻量模型加速解析, temperature=0.0 保证确定性输出
        result, usage = ark_client.chat_json_lite(messages, temperature=0.0, max_tokens=2048)
        logger.info(f"简历结构化完成, 字段: {list(result.keys())}, tokens={usage}")
        return result, usage

    def parse_job(self, text: str) -> tuple[dict[str, Any], dict[str, int]]:
        """用灵犀大模型将职位描述结构化, 返回 (结构化数据, token_usage)"""
        if not text.strip():
            raise ValueError("职位文本为空")
        truncated = text[:4000]
        messages = build_job_messages(truncated)
        # 使用 mini 轻量模型加速解析
        result, usage = ark_client.chat_json_lite(messages, temperature=0.0, max_tokens=1024)
        # 字段映射容错: 兼容灵犀返回的别名
        field_aliases = {
            "title": ["title", "job_title", "position", "职位名称", "职位"],
            "company": ["company", "company_name", "公司名称", "公司"],
            "department": ["department", "dept", "部门"],
            "job_type": ["job_type", "employment_type", "工作性质", "工作类型"],
            "salary_min": ["salary_min", "min_salary", "薪资下限", "最低薪资"],
            "salary_max": ["salary_max", "max_salary", "薪资上限", "最高薪资"],
            "work_city": ["work_city", "city", "location", "工作城市", "工作地点"],
            "experience_required": ["experience_required", "experience", "经验要求", "工作经验"],
            "education_required": ["education_required", "education", "学历要求", "学历"],
            "headcount": ["headcount", "hire_count", "招聘人数", "人数"],
            "description": ["description", "job_description", "职位描述", "岗位职责", "工作内容"],
            "requirements": ["requirements", "skills", "技能要求", "任职要求"],
        }
        normalized = {}
        for target, aliases in field_aliases.items():
            for alias in aliases:
                if alias in result and result[alias] is not None:
                    normalized[target] = result[alias]
                    break
            else:
                normalized[target] = None
        # 确保 requirements 是列表
        if normalized.get("requirements") and not isinstance(normalized["requirements"], list):
            normalized["requirements"] = []
        # 确保数值字段为整数
        for num_field in ["salary_min", "salary_max", "headcount"]:
            val = normalized.get(num_field)
            if val is not None:
                try:
                    normalized[num_field] = int(float(val))
                except (ValueError, TypeError):
                    normalized[num_field] = None
        filled = [k for k, v in normalized.items() if v is not None]
        logger.info(f"职位结构化完成, 有效字段 {len(filled)}/{len(normalized)}: {filled}, tokens={usage}")
        return normalized, usage


doc_parser = DocParser()
